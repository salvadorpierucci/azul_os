"""Generación de presupuesto en formato Word (.docx).

Estructura:
  - Word Cliente: formato plantilla (fecha, cliente, items con descripción,
    sin precios por ítem ni desglose de logística, solo total final,
    condiciones de contratación).
  - Word Completo: tabla con precios por ítem, desglose de logística,
    sección de fotos.
"""
import io
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD = RGBColor(0xC5, 0xA8, 0x80)
NAVY = RGBColor(0x3A, 0x58, 0x74)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x99, 0x99, 0x99)
COMPANY = "Azul Livings Luján"


def _fmt_money(val: float) -> str:
    return f"${int(round(val)):,}".replace(",", ".")


def _shade_cell(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _header_block(doc, ppto_data: dict, subtitle: str):
    p = doc.add_paragraph()
    run = p.add_run(COMPANY.upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.size = Pt(10)
    run.font.color.rgb = GOLD
    p.paragraph_format.space_after = Pt(8)

    info = doc.add_table(rows=3, cols=4)
    info.style = "Table Grid"
    info_data = [
        ("Cliente:", ppto_data.get("cliente_nombre") or "—", "Fecha evento:", ppto_data.get("fecha_evento") or "—"),
        ("Tipo:", ppto_data.get("tipo_evento") or "—", "Invitados:", str(ppto_data.get("cantidad_invitados") or "—")),
        ("Localidad:", ppto_data.get("localidad") or "—", "Traslado:", _fmt_money(ppto_data.get("costo_logistica") or 0)),
    ]
    for row_idx, (l1, v1, l2, v2) in enumerate(info_data):
        cells = info.rows[row_idx].cells
        for cell_idx, text in enumerate([l1, v1, l2, v2]):
            para = cells[cell_idx].paragraphs[0]
            run = para.add_run(text)
            if cell_idx % 2 == 0:
                run.bold = True
            run.font.size = Pt(10)


def _totales_block(doc, ppto_data: dict, show_mobiliario: bool = False):
    if show_mobiliario:
        tot_para = doc.add_paragraph()
        r = tot_para.add_run(f"Mobiliario: {_fmt_money(ppto_data.get('subtotal_mobiliario', 0))}")
        r.font.size = Pt(10)
        tot_para.paragraph_format.space_after = Pt(2)

    # Sección Logística: Traslado + Armado/Desarme
    costo_traslado = ppto_data.get("costo_logistica") or 0
    costo_armado = ppto_data.get("costo_armado") or 0
    if costo_traslado > 0 or costo_armado > 0:
        log_header = doc.add_paragraph()
        log_header.paragraph_format.space_before = Pt(4)
        log_header.paragraph_format.space_after = Pt(2)
        rh = log_header.add_run("Logística:")
        rh.bold = True
        rh.font.size = Pt(10)

        if costo_traslado > 0:
            sub_t = doc.add_paragraph()
            sub_t.paragraph_format.left_indent = Pt(14)
            sub_t.paragraph_format.space_after = Pt(1)
            r = sub_t.add_run(f"• Traslado: {_fmt_money(costo_traslado)}")
            r.font.size = Pt(10)

        if costo_armado > 0:
            sub_a = doc.add_paragraph()
            sub_a.paragraph_format.left_indent = Pt(14)
            sub_a.paragraph_format.space_after = Pt(1)
            r = sub_a.add_run(f"• Armado y desarme: {_fmt_money(costo_armado)}")
            r.font.size = Pt(10)

    # Descuento
    descuento = ppto_data.get("descuento") or 0
    if descuento > 0:
        desc_para = doc.add_paragraph()
        desc_para.paragraph_format.space_before = Pt(4)
        desc_para.paragraph_format.space_after = Pt(2)
        r = desc_para.add_run(f"Descuento: -{_fmt_money(descuento)}")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

    total_para = doc.add_paragraph()
    total_para.paragraph_format.space_before = Pt(6)
    r1 = total_para.add_run("TOTAL: ")
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = NAVY
    r2 = total_para.add_run(_fmt_money(ppto_data.get("total", 0)))
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = NAVY

    senia_para = doc.add_paragraph()
    senia_para.paragraph_format.space_before = Pt(6)
    run = senia_para.add_run("Seña del 50% para confirmar la reserva")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


def _totales_block_cliente(doc, ppto_data: dict):
    """Bloque de totales para el presupuesto CLIENTE: solo TOTAL, sin desglose."""
    total_para = doc.add_paragraph()
    total_para.paragraph_format.space_before = Pt(8)
    r1 = total_para.add_run("Importe total del servicio ")
    r1.font.size = Pt(12)
    r2 = total_para.add_run(_fmt_money(ppto_data.get("total", 0)))
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = NAVY


def _condiciones_block(doc):
    """Bloque de condiciones de contratación (formato plantilla)."""
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("CONDICIONES DE CONTRATACIÓN:")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run(
        "Reservando dentro de los primeros 3 días de recibido el "
        "presupuesto te congelamos el presupuesto!"
    )
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    condiciones = [
        "Con el 30% del valor total a alquilar, se reserva la fecha.",
        "La seña puede ser realizada mediante transferencia al alias "
        "azul.livings.lujan .El resto se abona en efectivo (actualizado, "
        "si no congeló presupuesto) el día del evento.",
        "En caso de postergar el evento, la seña será guardada durante "
        "un mes.",
        "El cambio de fecha quedará sujeto a la disponibilidad del "
        "mobiliario.",
        "En el caso de suspender definitivamente el evento, la seña del "
        "30% por el guardado de fecha, no tiene devolución.",
        "Todo lo contratado es en concepto de alquiler. En caso de haber "
        "roturas y/o faltantes de accesorios o mobiliario, el cliente "
        "deberá abonar el costo de la reposición.",
    ]
    for cond in condiciones:
        p = doc.add_paragraph()
        run = p.add_run(cond)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(7)


def _fotos_section(doc, lugares_raw: list, mob_fotos: dict):
    """Sección de fotos: nombre del mobiliario + foto grande, una por página."""
    # Recopilar items únicos que tienen foto
    items_con_foto = []
    vistos = set()
    for lugar in lugares_raw:
        for prod in lugar.get("productos", []):
            key = prod.get("catalogo_key", "") or prod.get("nombre", "")
            if key and key not in vistos:
                foto_path = mob_fotos.get(key)
                if foto_path and os.path.exists(foto_path):
                    items_con_foto.append((key, foto_path))
                    vistos.add(key)

    if not items_con_foto:
        return

    # Salto de página antes de la sección de fotos
    doc.add_page_break()

    # Título de la sección
    p = doc.add_paragraph()
    run = p.add_run("Fotos del mobiliario")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(12)

    for nombre, foto_path in items_con_foto:
        # Nombre del mobiliario
        p = doc.add_paragraph()
        run = p.add_run(nombre)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = NAVY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        # Foto grande centrada
        try:
            photo_para = doc.add_paragraph()
            photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = photo_para.add_run()
            run.add_picture(foto_path, width=Cm(14))
        except Exception:
            pass

        # Salto de página entre fotos (excepto la última)
        if (nombre, foto_path) != items_con_foto[-1]:
            doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# WORD CLIENTE — formato plantilla: sin precios, sin desglose,
# solo total final + condiciones de contratación
# ════════════════════════════════════════════════════════════════
def generate_word_cliente(ppto_data: dict, lugares_raw: list,
                          mob_fotos: dict = None) -> io.BytesIO:
    mob_fotos = mob_fotos or {}
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ─── FECHA (arriba a la derecha) ───
    fecha_display = _fmt_fecha_para_word(ppto_data.get("fecha_evento", ""))
    if fecha_display:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(fecha_display)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(8)

    # ─── DATOS DEL CLIENTE ───
    campos = [
        ("CLIENTE:", ppto_data.get("cliente_nombre") or ""),
        ("FECHA DEL EVENTO:", ppto_data.get("fecha_evento") or ""),
        ("LUGAR DEL EVENTO:", ppto_data.get("localidad") or ""),
        ("TIPO DE EVENTO:", ppto_data.get("tipo_evento") or ""),
    ]
    invitados = ppto_data.get("cantidad_invitados")
    if invitados:
        campos.append(("CANTIDAD DE INVITADOS:", str(invitados)))

    for label, valor in campos:
        p = doc.add_paragraph()
        run_label = p.add_run(label + " ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_valor = p.add_run(valor or "")
        run_valor.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # espacio

    # ─── MOBILIARIO (ordenado alfabéticamente, con descripción) ───
    # Recopilar todos los productos de todos los lugares
    todos_items = []
    for lugar in lugares_raw:
        for prod in lugar.get("productos", []):
            todos_items.append(prod)

    # Ordenar alfabéticamente por nombre
    todos_items.sort(key=lambda p: (p.get("catalogo_key", "") or p.get("nombre", "") or "").lower())

    if todos_items:
        p = doc.add_paragraph()
        run = p.add_run("Mobiliario:")
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

        for prod in todos_items:
            key = prod.get("catalogo_key", "") or prod.get("nombre", "")
            qty = prod.get("cantidad", 1)
            desc = prod.get("descripcion", "") or prod.get("notas", "")
            # Construir línea: "*48 sillas cross Back, en madera de roble, con almohadón."
            linea = f"*{qty} {key}"
            if desc:
                linea += f", {desc}"
            linea += "."
            p = doc.add_paragraph()
            run = p.add_run(linea)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(14)

    doc.add_paragraph()

    # ─── LOGÍSTICA (solo mención cualitativa, sin precios) ───
    localidad = ppto_data.get("localidad") or ""
    distancia = ppto_data.get("distancia_km")
    if localidad or distancia:
        p = doc.add_paragraph()
        run = p.add_run("Logística:")
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

        if localidad and distancia:
            txt_log = f"*Entrega y retiro por {localidad}."
        elif localidad:
            txt_log = f"*Entrega y retiro por {localidad}."
        else:
            txt_log = "*Entrega y retiro a coordinar."
        p = doc.add_paragraph()
        run = p.add_run(txt_log)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Pt(14)

    doc.add_paragraph()

    # ─── TOTAL ───
    _totales_block_cliente(doc, ppto_data)

    # ─── CONDICIONES DE CONTRATACIÓN ───
    _condiciones_block(doc)

    # ─── SECCIÓN DE FOTOS ───
    _fotos_section(doc, lugares_raw, mob_fotos)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════════════════════════
# WORD COMPLETO — con precios por item, desglose, luego sección de fotos
# ════════════════════════════════════════════════════════════════
def generate_word_completo(ppto_data: dict, lugares_raw: list,
                           mob_prices: dict, mob_fotos: dict = None) -> io.BytesIO:
    mob_fotos = mob_fotos or {}
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ─── HEADER + INFO ───
    _header_block(doc, ppto_data, "Presupuesto de Alquiler — Detalle Completo")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ─── TABLA DE ITEMS (sin fotos, con precios, ordenados alfabéticamente) ───
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, label in enumerate(["Item", "Cant.", "Precio Unit.", "Subtotal", "Descripción"]):
        para = hdr[i].paragraphs[0]
        run = para.add_run(label)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        _shade_cell(hdr[i], "3A5874")

    for lugar in lugares_raw:
        row = tbl.add_row()
        row.cells[1].merge(row.cells[2]).merge(row.cells[3]).merge(row.cells[4])
        cell = row.cells[0]
        para = cell.paragraphs[0]
        run = para.add_run(lugar.get("nombre", "General"))
        run.bold = True
        run.font.size = Pt(11)
        _shade_cell(cell, "F5F0E8")

        # Ordenar productos del lugar alfabéticamente
        productos = sorted(
            lugar.get("productos", []),
            key=lambda p: (p.get("catalogo_key", "") or p.get("nombre", "") or "").lower()
        )

        for prod in productos:
            key = prod.get("catalogo_key", "") or prod.get("nombre", "")
            qty = prod.get("cantidad", 1)
            # Usar precio guardado en el producto (calculado al guardar el presupuesto).
            # Si no existe (presupuesto viejo), caer a mob_prices por catalogo_key.
            precio_unit = prod.get("precio_unitario")
            if precio_unit is None or precio_unit == 0:
                precio_unit = mob_prices.get(key, 0)
            subtotal = prod.get("subtotal")
            if subtotal is None or subtotal == 0:
                subtotal = qty * precio_unit
            descripcion = prod.get("descripcion", "") or prod.get("notas", "") or ""

            row = tbl.add_row()
            row.cells[0].paragraphs[0].add_run(key).font.size = Pt(9)
            row.cells[1].paragraphs[0].add_run(str(qty)).font.size = Pt(9)
            row.cells[2].paragraphs[0].add_run(_fmt_money(precio_unit)).font.size = Pt(9)
            row.cells[3].paragraphs[0].add_run(_fmt_money(subtotal)).font.size = Pt(9)
            row.cells[4].paragraphs[0].add_run(descripcion).font.size = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ─── TOTALES ───
    _totales_block(doc, ppto_data, show_mobiliario=True)

    # ─── SECCIÓN DE FOTOS ───
    _fotos_section(doc, lugares_raw, mob_fotos)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _fmt_fecha_para_word(fecha_str: str) -> str:
    """Convierte '2026-07-03' → '3/7/26'."""
    if not fecha_str:
        return ""
    import re
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", str(fecha_str))
    if not m:
        return fecha_str
    return f"{int(m.group(3))}/{int(m.group(2))}/{m.group(1)[2:]}"
